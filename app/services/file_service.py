"""
Service pour la gestion des fichiers (upload, traitement, etc.).
"""

import os
import tempfile
import shutil
from typing import Optional, Dict, Any
from pathlib import Path
import logging
import asyncio
from datetime import datetime

import aiofiles
from fastapi import UploadFile, HTTPException
import magic
import PyPDF2
from docx import Document
from sqlalchemy.orm import Session

logger = logging.getLogger(__name__)

class FileService:
    """Service pour gérer les opérations sur les fichiers."""
    
    def __init__(self):
        """Initialise le service de fichiers."""
        self.upload_dir = Path("uploads")
        self.upload_dir.mkdir(exist_ok=True)
        self.allowed_extensions = {'.pdf', '.doc', '.docx', '.txt'}
        self.max_file_size = 10 * 1024 * 1024  # 10MB
    
    async def save_upload_file(self, file: UploadFile) -> str:
        """
        Sauvegarde un fichier uploadé.
        
        Args:
            file: Fichier uploadé
            
        Returns:
            Chemin du fichier sauvegardé
        """
        try:
            # Vérifier le type de fichier
            if not file.content_type:
                raise HTTPException(
                    status_code=400, 
                    detail="Type de fichier non déterminé"
                )
            
            # Vérifier l'extension
            file_extension = Path(file.filename).suffix.lower()
            if file_extension not in self.allowed_extensions:
                raise HTTPException(
                    status_code=400,
                    detail=f"Type de fichier non supporté. Extensions autorisées: {', '.join(self.allowed_extensions)}"
                )
            
            # Vérifier la taille du fichier
            contents = await file.read()
            if len(contents) > self.max_file_size:
                raise HTTPException(
                    status_code=400,
                    detail=f"Fichier trop volumineux. Taille maximale: {self.max_file_size // (1024*1024)}MB"
                )
            
            # Réinitialiser le curseur du fichier
            await file.seek(0)
            
            # Générer un nom de fichier unique
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_filename = f"{timestamp}_{file.filename}"
            file_path = self.upload_dir / unique_filename
            
            # Sauvegarder le fichier
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(contents)
            
            logger.info(f"Fichier sauvegardé: {file_path}")
            return str(file_path)
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Erreur lors de la sauvegarde du fichier: {e}")
            raise HTTPException(
                status_code=500,
                detail="Erreur lors de la sauvegarde du fichier"
            )
    
    async def extract_text_from_file(self, file_path: str) -> str:
        """
        Extrait le texte d'un fichier (PDF, DOC, DOCX).
        
        Args:
            file_path: Chemin du fichier
            
        Returns:
            Texte extrait du fichier
        """
        try:
            file_path = Path(file_path)
            file_extension = file_path.suffix.lower()
            
            if file_extension == '.pdf':
                return await self._extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                return await self._extract_text_from_docx(file_path)
            elif file_extension == '.doc':
                return await self._extract_text_from_doc(file_path)
            elif file_extension == '.txt':
                return await self._extract_text_from_txt(file_path)
            else:
                raise ValueError(f"Type de fichier non supporté: {file_extension}")
                
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction du texte: {e}")
            raise HTTPException(
                status_code=500,
                detail="Erreur lors de l'extraction du texte du fichier"
            )
    
    async def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extrait le texte d'un fichier PDF avec gestion robuste des erreurs."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Vérifier si le PDF est chiffré
                if pdf_reader.is_encrypted:
                    try:
                        pdf_reader.decrypt('')
                    except:
                        raise ValueError("Le PDF est protégé par mot de passe")
                
                # Extraire le texte de chaque page
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as page_error:
                        logger.warning(f"Erreur extraction page {page_num}: {page_error}")
                        continue
            
            # Vérifier que du texte a été extrait
            if not text.strip():
                raise ValueError("Aucun texte n'a pu être extrait du PDF. Le fichier pourrait être scanné ou contenir uniquement des images.")
            
            return text.strip()
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction du texte PDF: {e}")
            raise ValueError(f"Impossible d'extraire le texte du PDF: {str(e)}")
    
    async def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extrait le texte d'un fichier DOCX avec support des tableaux."""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extraire le texte des paragraphes
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extraire le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text += row_text + "\n"
            
            # Vérifier que du texte a été extrait
            if not text.strip():
                raise ValueError("Aucun texte n'a pu être extrait du document Word")
            
            return text.strip()
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction du texte DOCX: {e}")
            raise ValueError(f"Impossible d'extraire le texte du document Word: {str(e)}")
    
    async def _extract_text_from_doc(self, file_path: Path) -> str:
        """Extrait le texte d'un fichier DOC."""
        try:
            # Pour les fichiers .doc, on pourrait utiliser textract ou une conversion
            # Pour l'instant, retourner un message indiquant que le format n'est pas supporté
            logger.warning(f"Format .doc non supporté pour l'extraction de texte: {file_path}")
            return "[Format .doc - extraction de texte non supportée]"
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction du texte DOC: {e}")
            raise
    
    async def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extrait le texte d'un fichier TXT."""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                text = await file.read()
            return text.strip()
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction du texte TXT: {e}")
            raise
    
    async def cleanup_file(self, file_path: str) -> bool:
        """
        Nettoie un fichier temporaire.
        
        Args:
            file_path: Chemin du fichier à nettoyer
            
        Returns:
            True si le nettoyage a réussi
        """
        try:
            file_path = Path(file_path)
            if file_path.exists():
                file_path.unlink()
                logger.info(f"Fichier nettoyé: {file_path}")
                return True
            return False
            
        except Exception as e:
            logger.error(f"Erreur lors du nettoyage du fichier: {e}")
            return False
    
    async def export_jobs_data(
        self, 
        db: Session, 
        format: str = "csv"
    ) -> str:
        """
        Exporte les données des offres d'emploi.
        
        Args:
            db: Session de base de données
            format: Format d'export (csv, json, xlsx)
            
        Returns:
            Chemin du fichier exporté
        """
        try:
            if format not in ["csv", "json", "xlsx"]:
                raise ValueError(f"Format d'export non supporté: {format}")
            
            # Ici vous implémenteriez la logique d'export
            # Pour l'instant, retourner un exemple
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            export_filename = f"jobs_export_{timestamp}.{format}"
            export_path = self.upload_dir / "exports" / export_filename
            
            # Créer le répertoire d'exports s'il n'existe pas
            export_path.parent.mkdir(exist_ok=True)
            
            # Logique d'export simplifiée
            if format == "csv":
                await self._export_to_csv(db, export_path)
            elif format == "json":
                await self._export_to_json(db, export_path)
            elif format == "xlsx":
                await self._export_to_xlsx(db, export_path)
            
            logger.info(f"Données exportées: {export_path}")
            return str(export_path)
            
        except Exception as e:
            logger.error(f"Erreur lors de l'export des données: {e}")
            raise HTTPException(
                status_code=500,
                detail="Erreur lors de l'export des données"
            )
    
    async def _export_to_csv(self, db: Session, export_path: Path) -> None:
        """Exporte les données en format CSV."""
        # Implémentation de l'export CSV
        pass
    
    async def _export_to_json(self, db: Session, export_path: Path) -> None:
        """Exporte les données en format JSON."""
        # Implémentation de l'export JSON
        pass
    
    async def _export_to_xlsx(self, db: Session, export_path: Path) -> None:
        """Exporte les données en format Excel."""
        # Implémentation de l'export Excel
        pass
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """
        Obtient les informations sur un fichier.
        
        Args:
            file_path: Chemin du fichier
            
        Returns:
            Dictionnaire avec les informations du fichier
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"Fichier non trouvé: {file_path}")
            
            stat = file_path.stat()
            
            # Détecter le type MIME
            mime = magic.Magic(mime=True)
            mime_type = mime.from_file(str(file_path))
            
            return {
                "filename": file_path.name,
                "size": stat.st_size,
                "size_human": self._format_file_size(stat.st_size),
                "mime_type": mime_type,
                "extension": file_path.suffix.lower(),
                "created": datetime.fromtimestamp(stat.st_ctime),
                "modified": datetime.fromtimestamp(stat.st_mtime),
                "path": str(file_path)
            }
            
        except Exception as e:
            logger.error(f"Erreur lors de la récupération des infos du fichier: {e}")
            raise
    
    def _format_file_size(self, size: int) -> str:
        """Formate la taille du fichier en format lisible."""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024.0:
                return f"{size:.1f} {unit}"
            size /= 1024.0
        return f"{size:.1f} TB"
    
    async def validate_cv_file(self, file_path: str) -> Dict[str, Any]:
        """
        Valide et analyse un fichier CV.
        
        Args:
            file_path: Chemin du fichier CV
            
        Returns:
            Dictionnaire avec les informations de validation
        """
        try:
            file_info = self.get_file_info(file_path)
            text_content = await self.extract_text_from_file(file_path)
            
            # Analyse basique du contenu
            word_count = len(text_content.split())
            char_count = len(text_content)
            
            # Validation basique du contenu
            is_valid = (
                word_count > 50 and  # Au moins 50 mots
                char_count > 200 and  # Au moins 200 caractères
                file_info["size"] < self.max_file_size  # Taille valide
            )
            
            return {
                "is_valid": is_valid,
                "file_info": file_info,
                "content_analysis": {
                    "word_count": word_count,
                    "char_count": char_count,
                    "has_contact_info": any(keyword in text_content.lower() for keyword in ['email', 'téléphone', 'phone', 'contact']),
                    "has_experience": any(keyword in text_content.lower() for keyword in ['expérience', 'experience', 'emploi', 'job']),
                    "has_education": any(keyword in text_content.lower() for keyword in ['formation', 'education', 'diplôme', 'degree'])
                }
            }
            
        except Exception as e:
            logger.error(f"Erreur lors de la validation du CV: {e}")
            return {
                "is_valid": False,
                "error": str(e)
            }
