"""
Service de génération intelligente de documents (CV, Lettres, etc.)
- Gemini pour la rédaction anti-hallucination (données utilisateur uniquement)
- ReportLab pour l'export PDF professionnel
"""
import os
import io
import json
import logging
import uuid
from datetime import datetime
from typing import Optional, Dict, Any, List
from dataclasses import dataclass, field

import re

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
    Table, TableStyle, KeepTogether, Image as RLImage
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY, TA_RIGHT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.graphics.shapes import Drawing, Circle, Rect, String

from .llm_client import LLMClient

logger = logging.getLogger(__name__)

# ────────────────────────────────────────────────────────────
# PALETTE DE COULEURS
# ────────────────────────────────────────────────────────────
PRIMARY     = colors.HexColor("#0F766E")   # Turquoise profond / Sahel Teal
PRIMARY_DK  = colors.HexColor("#0B5650")   # Teal foncé (fonds pleins, contrastes)
ACCENT      = colors.HexColor("#0EA5E9")   # Cyan — accent ponctuel
TINT        = colors.HexColor("#E6F6F4")   # Teinte très claire (puces, tags, fonds doux)
TINT_STRONG = colors.HexColor("#CCEEEA")   # Teinte un peu plus marquée (fond des tags)
SIDEBAR_BG  = colors.HexColor("#F4FAF9")   # Fond colonne latérale (légère teinte teal)
LINE        = colors.HexColor("#DCE7E5")   # Filets / séparateurs
INK         = colors.HexColor("#0F172A")   # Quasi-noir — nom, titres forts
TEXT_DARK   = colors.HexColor("#1E293B")   # Texte principal Slate
TEXT_GREY   = colors.HexColor("#64748B")   # Texte secondaire Slate
WHITE       = colors.white

# Conservé pour compatibilité avec d'éventuels appels existants
LIGHT_BG = SIDEBAR_BG

OUTPUT_DIR = "app/static/generated"
os.makedirs(OUTPUT_DIR, exist_ok=True)

_MOIS_FR = [
    "janvier", "février", "mars", "avril", "mai", "juin",
    "juillet", "août", "septembre", "octobre", "novembre", "décembre",
]


def _format_date_fr(dt: datetime) -> str:
    """Formate une date en français (ex. 22 juin 2026) sans dépendre de la locale système."""
    return f"{dt.day} {_MOIS_FR[dt.month - 1]} {dt.year}"


@dataclass
class UserProfileData:
    """Données du profil candidat pour la génération."""
    first_name: str = ""
    last_name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    current_title: str = ""
    experience_years: int = 0
    education_level: str = ""
    skills: List[str] = field(default_factory=list)
    bio: str = ""
    linkedin: str = ""
    github: str = ""
    portfolio: str = ""
    languages: Optional[Any] = None
    experiences: Optional[Any] = None
    certifications: Optional[Any] = None
    avatar_url: Optional[str] = None
    interests: Optional[List[str]] = field(default_factory=list)


@dataclass
class GeneratedDocument:
    """Résultat de la génération."""
    content_text: str
    file_path: str
    file_name: str
    document_type: str
    docx_file_path: Optional[str] = None
    docx_file_name: Optional[str] = None


class DocumentGenerationService:
    """Génère des documents professionnels PDF via Gemini + ReportLab."""

    def __init__(self):
        self.llm = LLMClient()

    # ────────────────────────────────────────────────────────
    # API PUBLIQUE
    # ────────────────────────────────────────────────────────

    async def generate_cover_letter(
        self,
        profile: UserProfileData,
        job_title: str,
        company_name: str,
        job_description: str = "",
        tone: str = "professionnel",
    ) -> GeneratedDocument:
        """Génère une lettre de motivation sur mesure, strictement ancrée dans le profil."""
        content = await self._llm_cover_letter(profile, job_title, company_name, job_description, tone)
        
        # PDF
        file_path = self._build_letter_pdf(
            content=content,
            profile=profile,
            title=f"Lettre de motivation – {job_title} chez {company_name}",
            doc_type="cover_letter",
        )
        file_name = os.path.basename(file_path)
        
        # Word (.docx)
        docx_file_path = self._build_letter_docx(
            content=content,
            profile=profile,
            title=f"Lettre de motivation – {job_title} chez {company_name}",
            doc_type="cover_letter",
        )
        docx_file_name = os.path.basename(docx_file_path)
        
        return GeneratedDocument(
            content_text=content,
            file_path=file_path,
            file_name=file_name,
            document_type="cover_letter",
            docx_file_path=docx_file_path,
            docx_file_name=docx_file_name
        )

    async def generate_cv(
        self,
        profile: UserProfileData,
        target_job: str = "",
    ) -> GeneratedDocument:
        """Génère un CV PDF et Word structuré à partir du profil réel."""
        summary = await self._llm_cv_summary(profile, target_job)
        profile.bio = summary  # on injecte l'accroche IA
        
        # PDF
        file_path = self._build_cv_pdf(profile, target_job)
        file_name = os.path.basename(file_path)
        
        # Word (.docx)
        docx_file_path = self._build_cv_docx(profile, target_job)
        docx_file_name = os.path.basename(docx_file_path)
        
        return GeneratedDocument(
            content_text=summary,
            file_path=file_path,
            file_name=file_name,
            document_type="cv",
            docx_file_path=docx_file_path,
            docx_file_name=docx_file_name
        )

    async def generate_other_letter(
        self,
        profile: UserProfileData,
        letter_type: str,
        context: Dict[str, str],
    ) -> GeneratedDocument:
        """Génère une lettre administrative (démission, relance, stage, recommandation)."""
        content = await self._llm_other_letter(profile, letter_type, context)
        
        # PDF
        file_path = self._build_letter_pdf(
            content=content,
            profile=profile,
            title=self._letter_type_label(letter_type),
            doc_type=letter_type,
        )
        file_name = os.path.basename(file_path)
        
        # Word (.docx)
        docx_file_path = self._build_letter_docx(
            content=content,
            profile=profile,
            title=self._letter_type_label(letter_type),
            doc_type=letter_type,
        )
        docx_file_name = os.path.basename(docx_file_path)
        
        return GeneratedDocument(
            content_text=content,
            file_path=file_path,
            file_name=file_name,
            document_type=letter_type,
            docx_file_path=docx_file_path,
            docx_file_name=docx_file_name
        )

    # ────────────────────────────────────────────────────────
    # GÉNÉRATION LLM (Gemini) — ANTI-HALLUCINATION
    # ────────────────────────────────────────────────────────

    async def _llm_cover_letter(
        self,
        profile: UserProfileData,
        job_title: str,
        company_name: str,
        job_description: str,
        tone: str,
    ) -> str:
        skills_str = ", ".join(profile.skills[:12]) if profile.skills else "Non précisé"
        exp_list = json.dumps(profile.experiences, ensure_ascii=False)[:800] if profile.experiences else "Non précisé"

        system = (
            "Tu es un expert RH sénégalais rédigeant des lettres de motivation professionnelles. "
            "RÈGLE ABSOLUE : n'invente AUCUNE information. Utilise UNIQUEMENT les données fournies. "
            "Si une donnée est manquante, construis une phrase générale sans mentir. "
            f"Ton : {tone}. Langue : français. Format : texte brut, sans markdown."
        )

        user = f"""Rédige une lettre de motivation complète (3 paragraphes) pour :
CANDIDAT :
- Nom : {profile.first_name} {profile.last_name}
- Poste actuel : {profile.current_title or 'Non précisé'}
- Expérience : {profile.experience_years} ans
- Compétences : {skills_str}
- Expériences détaillées : {exp_list}
- Bio : {profile.bio or 'Non fournie'}

OFFRE :
- Poste visé : {job_title}
- Entreprise : {company_name}
- Description (extrait) : {job_description[:500] if job_description else 'Non fournie'}

Structure :
Paragraphe 1 – Accroche + intérêt pour le poste (basé sur les données réelles)
Paragraphe 2 – Compétences et expériences pertinentes (UNIQUEMENT celles listées ci-dessus)
Paragraphe 3 – Motivation pour l'entreprise + formule de politesse

Commence directement par "Madame, Monsieur,"
"""
        return await self.llm.generate_response(system, user, temperature=0.4)

    async def _llm_cv_summary(self, profile: UserProfileData, target_job: str) -> str:
        skills_str = ", ".join(profile.skills[:10]) if profile.skills else "Non précisé"
        system = (
            "Tu es un expert en rédaction de CV. "
            "RÈGLE ABSOLUE : 2-3 phrases max, basées UNIQUEMENT sur les données fournies. "
            "Pas de markdown. Pas d'inventions."
        )
        user = f"""Rédige une accroche de profil percutante pour ce CV :
- Prénom/Nom : {profile.first_name} {profile.last_name}
- Titre : {profile.current_title or 'Professionnel en recherche'}
- Expérience : {profile.experience_years} ans
- Compétences clés : {skills_str}
- Poste visé : {target_job or 'Non précisé'}
- Bio existante : {profile.bio or 'Aucune'}
"""
        return await self.llm.generate_response(system, user, temperature=0.3)

    async def _llm_other_letter(
        self,
        profile: UserProfileData,
        letter_type: str,
        context: Dict[str, str],
    ) -> str:
        templates = {
            "resignation":      "une lettre de démission professionnelle et respectueuse",
            "internship_request": "une demande de stage formelle",
            "follow_up":        "une lettre de relance après candidature",
            "recommendation_request": "une demande de lettre de recommandation",
        }
        desc = templates.get(letter_type, "une lettre professionnelle")
        ctx_str = "\n".join(f"- {k}: {v}" for k, v in context.items())

        system = (
            f"Tu es expert en rédaction administrative. Rédige {desc}. "
            "RÈGLE : données fournies uniquement, pas d'invention. Texte brut, sans markdown."
        )
        user = f"""Candidat : {profile.first_name} {profile.last_name}
Titre actuel : {profile.current_title or 'Non précisé'}
Contexte supplémentaire :
{ctx_str}

Rédige la lettre complète avec formules de politesse adaptées.
"""
        return await self.llm.generate_response(system, user, temperature=0.3)

    # ────────────────────────────────────────────────────────
    # GÉNÉRATION PDF — LETTRE
    # ────────────────────────────────────────────────────────

    def _build_letter_pdf(
        self,
        content: str,
        profile: UserProfileData,
        title: str,
        doc_type: str,
    ) -> str:
        file_name = f"{doc_type}_{uuid.uuid4().hex[:8]}.pdf"
        file_path = os.path.join(OUTPUT_DIR, file_name)

        doc = SimpleDocTemplate(
            file_path, pagesize=A4,
            leftMargin=2.5*cm, rightMargin=2.5*cm,
            topMargin=2.5*cm, bottomMargin=2.5*cm,
        )

        styles = self._get_styles()
        story = []

        # ── En-tête candidat & destinataire ───────────────
        header_left = [
            Paragraph(f"<b>{profile.first_name} {profile.last_name}</b>", styles["body_normal"]),
            Paragraph(f"Tél : {profile.phone}", styles["body_normal"]) if profile.phone else "",
            Paragraph(f"Email : {profile.email}", styles["body_normal"]) if profile.email else "",
            Paragraph(f"Adresse : {profile.location}", styles["body_normal"]) if profile.location else "",
        ]

        recipient_name = "Responsable du Recrutement"
        company_lbl = ""
        if "chez" in title:
            parts = title.split("chez")
            company_lbl = parts[-1].strip()

        header_right = [
            Paragraph(f"Dakar, le {_format_date_fr(datetime.now())}", styles["date"]),
            Spacer(1, 15),
            Paragraph(f"<b>À l'attention du</b><br/><b>{recipient_name}</b>", styles["body_normal"]),
        ]
        if company_lbl:
            header_right.append(Paragraph(f"<i>{company_lbl}</i>", styles["body_normal"]))

        header_table = Table([[header_left, header_right]], colWidths=[10 * cm, 6 * cm])
        header_table.setStyle(TableStyle([
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("LEFTPADDING", (0,0), (-1,-1), 0),
            ("RIGHTPADDING", (0,0), (-1,-1), 0),
            ("TOPPADDING", (0,0), (-1,-1), 0),
            ("BOTTOMPADDING", (0,0), (-1,-1), 0),
        ]))
        story.append(header_table)
        story.append(Spacer(1, 1.2 * cm))

        # ── Titre lettre / Objet ──────────────────────────
        story.append(Paragraph(f"<b>Objet : {title}</b>", styles["letter_object"]))
        story.append(Spacer(1, 0.8 * cm))

        # ── Corps de la lettre ────────────────────────────
        for para in content.split("\n"):
            para = para.strip()
            if para:
                story.append(Paragraph(para, styles["body_justify"]))
                story.append(Spacer(1, 0.3 * cm))

        story.append(Spacer(1, 1.2 * cm))
        
        # Signature
        sig_data = [
            ["", Paragraph("<b>Le Candidat,</b>", styles["body_normal"])],
            ["", Paragraph(f"<b>{profile.first_name} {profile.last_name}</b>", styles["body_normal"])]
        ]
        sig_table = Table(sig_data, colWidths=[10 * cm, 6 * cm])
        sig_table.setStyle(TableStyle([
            ("ALIGN", (1,0), (1,-1), "LEFT"),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("LEFTPADDING", (0,0), (-1,-1), 0),
            ("RIGHTPADDING", (0,0), (-1,-1), 0),
        ]))
        story.append(sig_table)

        doc.build(story)
        logger.info(f"PDF lettre générée : {file_path}")
        return file_path

    def _build_letter_docx(
        self,
        content: str,
        profile: UserProfileData,
        title: str,
        doc_type: str,
    ) -> str:
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import qn, nsdecls

        file_name = f"{doc_type}_{uuid.uuid4().hex[:8]}.docx"
        file_path = os.path.join(OUTPUT_DIR, file_name)

        doc = docx.Document()

        # Marges à 1 pouce (2.5 cm)
        margin = Inches(1)
        for section in doc.sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        # Couleurs
        primary_color = RGBColor(13, 148, 136)     # #0D9488
        text_dark = RGBColor(30, 41, 59)          # #1E293B
        text_grey = RGBColor(71, 85, 105)         # #475569

        # Police de base
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10.5)
        font.color.rgb = text_dark

        # En-tête expéditeur & destinataire
        table = doc.add_table(rows=1, cols=2)
        tblPr = table._tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{border_name}')
            b.set(qn('w:val'), 'none')
            tblBorders.append(b)
        tblPr.append(tblBorders)

        col_widths = [Inches(3.8), Inches(2.7)]
        for row in table.rows:
            row.cells[0].width = col_widths[0]
            row.cells[1].width = col_widths[1]

        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]

        # Gauche : Expéditeur
        p_left = left_cell.paragraphs[0]
        p_left.paragraph_format.line_spacing = 1.15
        run_name = p_left.add_run(f"{profile.first_name} {profile.last_name}\n")
        run_name.bold = True
        run_name.font.size = Pt(11)
        run_name.font.color.rgb = primary_color
        
        contacts = []
        if profile.phone: contacts.append(f"Tél : {profile.phone}")
        if profile.email: contacts.append(f"Email : {profile.email}")
        if profile.location: contacts.append(f"Adresse : {profile.location}")
        
        run_contacts = p_left.add_run("\n".join(contacts))
        run_contacts.font.size = Pt(9.5)
        run_contacts.font.color.rgb = text_dark

        # Droite : Date et Destinataire
        p_right = right_cell.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_right.paragraph_format.line_spacing = 1.15
        run_date = p_right.add_run(f"Dakar, le {_format_date_fr(datetime.now())}\n\n")
        run_date.font.size = Pt(9.5)
        run_date.font.color.rgb = text_grey
        
        recipient_name = "Responsable du Recrutement"
        company_lbl = ""
        if "chez" in title:
            parts = title.split("chez")
            company_lbl = parts[-1].strip()
            
        run_to = p_right.add_run("À l'attention du\n")
        run_to.bold = True
        run_to.font.size = Pt(9.5)
        run_to.font.color.rgb = text_dark
        
        run_rec = p_right.add_run(f"{recipient_name}\n")
        run_rec.bold = True
        run_rec.font.size = Pt(9.5)
        run_rec.font.color.rgb = text_dark
        
        if company_lbl:
            run_comp = p_right.add_run(company_lbl)
            run_comp.italic = True
            run_comp.font.size = Pt(9.5)
            run_comp.font.color.rgb = text_dark

        # Espace après le tableau d'en-tête
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(12)

        # Objet
        p_obj = doc.add_paragraph()
        p_obj.paragraph_format.space_after = Pt(18)
        run_obj = p_obj.add_run(f"Objet : {title}")
        run_obj.bold = True
        run_obj.font.size = Pt(11)
        run_obj.font.color.rgb = primary_color

        # Corps
        for para in content.split("\n"):
            para = para.strip()
            if para:
                p_body = doc.add_paragraph()
                p_body.paragraph_format.space_after = Pt(8)
                p_body.paragraph_format.line_spacing = 1.15
                p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run_body = p_body.add_run(para)
                run_body.font.size = Pt(10.5)

        # Signature
        p_sig = doc.add_paragraph()
        p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_sig.paragraph_format.space_before = Pt(24)
        run_sig_lbl = p_sig.add_run("Le Candidat,\n\n")
        run_sig_lbl.bold = True
        run_sig_lbl.font.size = Pt(10.5)
        run_sig_lbl.font.color.rgb = text_dark
        
        run_sig = p_sig.add_run(f"{profile.first_name} {profile.last_name}")
        run_sig.bold = True
        run_sig.font.size = Pt(10.5)
        run_sig.font.color.rgb = primary_color

        doc.save(file_path)
        logger.info(f"DOCX lettre générée : {file_path}")
        return file_path

    def _build_cv_docx(
        self,
        profile: UserProfileData,
        target_job: str = "",
    ) -> str:
        import docx
        from docx.shared import Inches, Pt, RGBColor, Mm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import qn, nsdecls

        file_name = f"cv_{uuid.uuid4().hex[:8]}.docx"
        file_path = os.path.join(OUTPUT_DIR, file_name)

        doc = docx.Document()

        # Format A4 + marges étroites à 0.5 pouce
        margin = Inches(0.5)
        for section in doc.sections:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        # Couleurs (alignées sur la palette PDF)
        teal_color  = RGBColor(0x0F, 0x76, 0x6E)   # PRIMARY  #0F766E
        teal_dark   = RGBColor(0x0B, 0x56, 0x50)   # PRIMARY_DK
        ink_color   = RGBColor(0x0F, 0x17, 0x2A)   # INK
        text_dark   = RGBColor(0x1E, 0x29, 0x3B)   # TEXT_DARK
        text_grey   = RGBColor(0x64, 0x74, 0x8B)   # TEXT_GREY
        line_grey   = "DCE7E5"                      # LINE
        sidebar_bg  = "F4FAF9"                       # SIDEBAR_BG
        chip_bg     = "CCEEEA"                        # TINT_STRONG

        # Style de base
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(9.5)
        normal_style.font.color.rgb = text_dark

        # Helpers XML internes
        def set_cell_shading(cell, color_hex):
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
            cell._tc.get_or_add_tcPr().append(shd)

        def set_table_fixed_layout(table, total_width):
            """Empêche Word de redimensionner le tableau selon son contenu (cause de débordement)."""
            table.autofit = False
            table.allow_autofit = False
            tblPr = table._tbl.tblPr
            layout = OxmlElement('w:tblLayout')
            layout.set(qn('w:type'), 'fixed')
            tblPr.append(layout)
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:type'), 'dxa')
            tblW.set(qn('w:w'), str(int(total_width / Inches(1) * 1440)))
            tblPr.append(tblW)

        def set_cell_left_border(cell, color_hex="DCE7E5", size="12"):
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), str(size))
            left.set(qn('w:space'), '0')
            left.set(qn('w:color'), color_hex)
            tcBorders.append(left)
            for side in ('top', 'bottom', 'right'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'none')
                tcBorders.append(b)
            tcPr.append(tcBorders)

        def remove_table_borders(table):
            tblPr = table._tbl.tblPr
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                b = OxmlElement(f'w:{border_name}')
                b.set(qn('w:val'), 'none')
                tblBorders.append(b)
            tblPr.append(tblBorders)

        # ── Tableau Principal à 2 colonnes ────────────────
        outer_table = doc.add_table(rows=1, cols=2)
        remove_table_borders(outer_table)

        col_widths = [Inches(2.3), Inches(4.85)]
        set_table_fixed_layout(outer_table, col_widths[0] + col_widths[1])
        outer_table.columns[0].width = col_widths[0]
        outer_table.columns[1].width = col_widths[1]
        for row in outer_table.rows:
            row.cells[0].width = col_widths[0]
            row.cells[1].width = col_widths[1]

        left_cell = outer_table.rows[0].cells[0]
        right_cell = outer_table.rows[0].cells[1]
        
        set_cell_shading(left_cell, sidebar_bg)
        set_cell_shading(right_cell, "FFFFFF")

        # ── Colonne Gauche (Sidebar) ──────────────────────
        # Photo (si disponible) ou monogramme / initiales — image circulaire (fiable, sans relief de fond carré)
        initials = ""
        if profile.first_name: initials += profile.first_name[0]
        if profile.last_name: initials += profile.last_name[0]
        if not initials: initials = "CV"

        p_mono = left_cell.paragraphs[0]
        p_mono.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_mono.paragraph_format.space_before = Pt(12)
        p_mono.paragraph_format.space_after = Pt(14)

        avatar_path = None
        if profile.avatar_url:
            try:
                import requests
                import tempfile
                r = requests.get(profile.avatar_url, timeout=5)
                if r.status_code == 200:
                    temp_photo = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
                    temp_photo.write(r.content)
                    temp_photo.close()
                    avatar_path = self._circular_image_path(temp_photo.name)
            except Exception as e:
                logger.warning(f"Impossible de charger la photo pour le Word : {e}")

        if not avatar_path:
            avatar_path = self._circular_monogram_path(initials)

        if avatar_path:
            p_mono.add_run().add_picture(avatar_path, width=Inches(1.15), height=Inches(1.15))
        else:
            run_mono = p_mono.add_run(f" {initials.upper()} ")
            run_mono.bold = True
            run_mono.font.size = Pt(22)
            run_mono.font.color.rgb = RGBColor(255, 255, 255)
            shd_mono = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0F766E"/>')
            run_mono._r.get_or_add_rPr().append(shd_mono)

        def add_rule(cell, space_before=8, space_after=10):
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(space_after)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), line_grey)
            pBdr.append(bottom)
            pPr.append(pBdr)
            return p

        # Coordonnées — libellés textuels (fiables, sans dépendance à une police emoji)
        def add_left_contact(label, text):
            p = left_cell.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            run_lbl = p.add_run(f"{label}\n")
            run_lbl.bold = True
            run_lbl.font.color.rgb = teal_color
            run_lbl.font.size = Pt(7.2)
            run_val = p.add_run(text)
            run_val.font.size = Pt(8.8)
            run_val.font.color.rgb = text_dark

        if profile.location: add_left_contact("LIEU", profile.location)
        if profile.phone: add_left_contact("TÉL", profile.phone)
        if profile.email: add_left_contact("EMAIL", profile.email)
        if profile.linkedin: add_left_contact("LINKEDIN", profile.linkedin)
        if profile.github: add_left_contact("GITHUB", profile.github)
        if profile.portfolio: add_left_contact("PORTFOLIO", profile.portfolio)

        add_rule(left_cell)

        # Compétences — étiquettes mises en valeur (fond teinté derrière le texte)
        if profile.skills:
            p_skills_lbl = left_cell.add_paragraph()
            p_skills_lbl.paragraph_format.space_after = Pt(6)
            run_lbl = p_skills_lbl.add_run("COMPÉTENCES")
            run_lbl.bold = True
            run_lbl.font.color.rgb = teal_color
            run_lbl.font.size = Pt(9.5)

            for skill in profile.skills:
                p_skill = left_cell.add_paragraph()
                p_skill.paragraph_format.space_after = Pt(5)
                run_chip = p_skill.add_run(f" {skill} ")
                run_chip.bold = True
                run_chip.font.size = Pt(8.3)
                run_chip.font.color.rgb = teal_dark
                shd_chip = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{chip_bg}"/>')
                run_chip._r.get_or_add_rPr().append(shd_chip)

            add_rule(left_cell)

        # Langues — un item structuré par langue, avec barre de niveau (colonne latérale)
        if profile.languages:
            p_lang_lbl = left_cell.add_paragraph()
            p_lang_lbl.paragraph_format.space_after = Pt(6)
            run_lbl = p_lang_lbl.add_run("LANGUES")
            run_lbl.bold = True
            run_lbl.font.color.rgb = teal_color
            run_lbl.font.size = Pt(9.5)

            langs = profile.languages if isinstance(profile.languages, list) else list(profile.languages.items())
            for lang in langs:
                if isinstance(lang, (list, tuple)):
                    label, level_str = lang[0], lang[1]
                else:
                    label, level_str = str(lang), "Moyen"

                lower_lvl = level_str.lower()
                if "maternelle" in lower_lvl or "bilingue" in lower_lvl or "parfait" in lower_lvl:
                    level_val = 1.0
                elif "courant" in lower_lvl or "avance" in lower_lvl or "avancé" in lower_lvl:
                    level_val = 0.8
                elif "intermediaire" in lower_lvl or "intermédiaire" in lower_lvl or "moyen" in lower_lvl:
                    level_val = 0.6
                else:
                    level_val = 0.3
                level_val = min(max(level_val, 0.06), 0.94)

                p_lname = left_cell.add_paragraph()
                p_lname.paragraph_format.space_after = Pt(0)
                run_lname = p_lname.add_run(label)
                run_lname.bold = True
                run_lname.font.size = Pt(9)
                run_lname.font.color.rgb = ink_color

                p_llvl = left_cell.add_paragraph()
                p_llvl.paragraph_format.space_after = Pt(3)
                run_llvl = p_llvl.add_run(level_str)
                run_llvl.font.size = Pt(7.8)
                run_llvl.font.color.rgb = text_grey

                side_bar_width = Inches(2.0)
                bar_table = left_cell.add_table(rows=1, cols=2)
                remove_table_borders(bar_table)
                for c in bar_table.rows[0].cells:
                    tcPr = c._tc.get_or_add_tcPr()
                    tcMar = OxmlElement('w:tcMar')
                    for side in ('top', 'bottom', 'left', 'right'):
                        m = OxmlElement(f'w:{side}')
                        m.set(qn('w:w'), '0')
                        m.set(qn('w:type'), 'dxa')
                        tcMar.append(m)
                    tcPr.append(tcMar)
                fill_w = Inches(2.0 * level_val)
                track_w = Inches(2.0 * (1 - level_val))
                set_table_fixed_layout(bar_table, side_bar_width)
                bar_table.columns[0].width = fill_w
                bar_table.columns[1].width = track_w
                bar_table.rows[0].cells[0].width = fill_w
                bar_table.rows[0].cells[1].width = track_w
                set_cell_shading(bar_table.rows[0].cells[0], "0F766E")
                set_cell_shading(bar_table.rows[0].cells[1], "DCE7E5")
                for c in bar_table.rows[0].cells:
                    c.paragraphs[0].paragraph_format.space_after = Pt(0)
                    c.paragraphs[0].paragraph_format.space_before = Pt(0)
                    c.paragraphs[0].paragraph_format.line_spacing = 1.0
                    run_pad = c.paragraphs[0].add_run(" ")
                    run_pad.font.size = Pt(2)

                p_lgap = left_cell.add_paragraph()
                p_lgap.paragraph_format.space_after = Pt(8)

            add_rule(left_cell)

        # Centres d'intérêt / Certifications — étiquettes mises en valeur (colonne latérale)
        interests_list_left = []
        if hasattr(profile, 'interests') and profile.interests:
            if isinstance(profile.interests, list):
                interests_list_left.extend(profile.interests)
            else:
                interests_list_left.append(str(profile.interests))

        left_title_text = "CENTRES D'INTÉRÊT"
        if not interests_list_left and hasattr(profile, 'certifications') and profile.certifications:
            left_title_text = "CERTIFICATIONS"
            certs = profile.certifications if isinstance(profile.certifications, list) else [profile.certifications]
            for cert in certs:
                label = cert.get("name", str(cert)) if isinstance(cert, dict) else str(cert)
                interests_list_left.append(label)

        if interests_list_left:
            p_int_lbl = left_cell.add_paragraph()
            p_int_lbl.paragraph_format.space_after = Pt(6)
            run_lbl = p_int_lbl.add_run(left_title_text)
            run_lbl.bold = True
            run_lbl.font.color.rgb = teal_color
            run_lbl.font.size = Pt(9.5)

            for item in interests_list_left:
                p_item = left_cell.add_paragraph()
                p_item.paragraph_format.space_after = Pt(5)
                run_chip = p_item.add_run(f" {item} ")
                run_chip.bold = True
                run_chip.font.size = Pt(8.3)
                run_chip.font.color.rgb = teal_dark
                shd_chip = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{chip_bg}"/>')
                run_chip._r.get_or_add_rPr().append(shd_chip)

        # ── Colonne Droite (Main Area) ────────────────────
        # Nom et titre
        p_name = right_cell.paragraphs[0]
        p_name.paragraph_format.space_before = Pt(8)
        p_name.paragraph_format.space_after = Pt(2)
        run_name = p_name.add_run(f"{profile.first_name} {profile.last_name}")
        run_name.bold = True
        run_name.font.size = Pt(22)
        run_name.font.color.rgb = ink_color

        p_title = right_cell.add_paragraph()
        p_title.paragraph_format.space_after = Pt(16)
        title_txt = target_job or profile.current_title or "Professionnel"
        run_title = p_title.add_run(title_txt.upper())
        run_title.bold = True
        run_title.font.size = Pt(11.5)
        run_title.font.color.rgb = teal_color

        # Profil professionnel (résumé) — colonne principale, sous le titre
        if profile.bio:
            p_profil_lbl = right_cell.add_paragraph()
            p_profil_lbl.paragraph_format.space_after = Pt(6)
            pPr = p_profil_lbl._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '2')
            bottom.set(qn('w:color'), line_grey)
            pBdr.append(bottom)
            pPr.append(pBdr)
            run_profil_lbl = p_profil_lbl.add_run("PROFIL PROFESSIONNEL")
            run_profil_lbl.bold = True
            run_profil_lbl.font.color.rgb = ink_color
            run_profil_lbl.font.size = Pt(11)

            p_profil = right_cell.add_paragraph()
            p_profil.paragraph_format.space_after = Pt(14)
            p_profil.paragraph_format.line_spacing = 1.2
            p_profil.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_profil = p_profil.add_run(profile.bio)
            run_profil.font.size = Pt(9.5)
            run_profil.font.color.rgb = text_dark

        # Timeline en tableau imbriqué (mise en page fixe pour éviter tout débordement)
        timeline_table = right_cell.add_table(rows=0, cols=2)
        remove_table_borders(timeline_table)

        time_widths = [Inches(0.3), Inches(4.45)]
        set_table_fixed_layout(timeline_table, time_widths[0] + time_widths[1])
        timeline_table.columns[0].width = time_widths[0]
        timeline_table.columns[1].width = time_widths[1]

        def add_timeline_section(title_text):
            row = timeline_table.add_row()
            row.cells[0].width = time_widths[0]
            row.cells[1].width = time_widths[1]

            p0 = row.cells[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run0 = p0.add_run("•")
            run0.bold = True
            run0.font.color.rgb = teal_color
            run0.font.size = Pt(16)

            p1 = row.cells[1].paragraphs[0]
            p1.paragraph_format.space_after = Pt(6)
            pPr = p1._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '2')
            bottom.set(qn('w:color'), line_grey)
            pBdr.append(bottom)
            pPr.append(pBdr)
            run1 = p1.add_run(title_text.upper())
            run1.bold = True
            run1.font.color.rgb = ink_color
            run1.font.size = Pt(11)

        def add_timeline_item(bold_title, italic_meta, period_text="", bullets=None):
            row = timeline_table.add_row()
            row.cells[0].width = time_widths[0]
            row.cells[1].width = time_widths[1]

            set_cell_left_border(row.cells[1], color_hex=line_grey)

            p1 = row.cells[1].paragraphs[0]
            p1.paragraph_format.space_after = Pt(1)
            if period_text:
                tab_stops = p1.paragraph_format.tab_stops
                tab_stops.add_tab_stop(time_widths[1] - Inches(0.05), WD_TAB_ALIGNMENT.RIGHT)
            run_title = p1.add_run(bold_title)
            run_title.bold = True
            run_title.font.size = Pt(10)
            run_title.font.color.rgb = ink_color
            if period_text:
                run_period = p1.add_run(f"\t{period_text}")
                run_period.bold = True
                run_period.font.size = Pt(8.3)
                run_period.font.color.rgb = teal_color

            if italic_meta:
                p_meta = row.cells[1].add_paragraph()
                p_meta.paragraph_format.space_after = Pt(4)
                run_meta = p_meta.add_run(italic_meta)
                run_meta.italic = True
                run_meta.font.size = Pt(8.7)
                run_meta.font.color.rgb = text_grey

            for b in (bullets or []):
                p_desc = row.cells[1].add_paragraph()
                p_desc.paragraph_format.space_after = Pt(3)
                p_desc.paragraph_format.line_spacing = 1.15
                p_desc.paragraph_format.left_indent = Inches(0.12)
                p_desc.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run_dash = p_desc.add_run("– ")
                run_dash.font.color.rgb = teal_color
                run_dash.font.size = Pt(8.7)
                run_desc = p_desc.add_run(b)
                run_desc.font.size = Pt(8.7)
                run_desc.font.color.rgb = text_dark

            p_end = row.cells[1].add_paragraph()
            p_end.paragraph_format.space_after = Pt(4)

        # 1. Experiences
        if profile.experiences:
            add_timeline_section("Parcours Professionnel")
            exps = profile.experiences if isinstance(profile.experiences, list) else [profile.experiences]
            for exp in exps:
                if isinstance(exp, dict):
                    job_title = exp.get("title", exp.get("poste", ""))
                    company   = exp.get("company", exp.get("entreprise", ""))
                    period    = exp.get("period", exp.get("periode", ""))
                    desc      = exp.get("description", exp.get("desc", ""))
                    add_timeline_item(job_title, company, period, self._split_sentences(desc))
                else:
                    add_timeline_item(str(exp), "")

        # 2. Formation (regroupée par bloc séparé par une ligne vide : diplôme + établissement)
        if profile.education_level:
            add_timeline_section("Formation")
            edu_blocks = [b for b in profile.education_level.split("\n\n") if b.strip()]
            if not edu_blocks:
                edu_blocks = [profile.education_level]
            for block in edu_blocks:
                lines = [l.strip() for l in block.split("\n") if l.strip()]
                if not lines:
                    continue
                degree = lines[0]
                school = " — ".join(lines[1:]) if len(lines) > 1 else ""
                add_timeline_item(degree, school)

        doc.save(file_path)
        logger.info(f"DOCX CV généré : {file_path}")
        return file_path

    # ────────────────────────────────────────────────────────
    # PETITS UTILITAIRES DE MISE EN FORME
    # ────────────────────────────────────────────────────────

    @staticmethod
    def _split_sentences(text: str) -> List[str]:
        """Découpe une description en phrases courtes pour un affichage en puces."""
        if not text:
            return []
        text = text.strip()
        # Découpage sur les fins de phrase (. ! ?) suivies d'une majuscule/espace
        parts = re.split(r'(?<=[.!?])\s+(?=[A-ZÀ-Ý])', text)
        parts = [p.strip() for p in parts if p.strip()]
        return parts if parts else [text]

    @staticmethod
    def _dot(color, diameter=6) -> Drawing:
        """Petit disque plein dessiné en vectoriel (fiable, indépendant des polices)."""
        d = Drawing(diameter, diameter)
        d.add(Circle(diameter / 2, diameter / 2, diameter / 2, fillColor=color, strokeColor=None))
        return d

    @staticmethod
    def _square_marker(color, size=8, radius=2) -> Drawing:
        """Petit carré arrondi plein, utilisé comme repère de section."""
        d = Drawing(size, size)
        d.add(Rect(0, 0, size, size, rx=radius, ry=radius, fillColor=color, strokeColor=None))
        return d

    def _section_header(self, text: str, styles) -> Table:
        """Titre de section (colonne principale) : carré accent + libellé + filet inférieur."""
        row = [self._square_marker(PRIMARY, size=9, radius=2), Paragraph(text.upper(), styles["cv_section_right"])]
        t = Table([row], colWidths=[0.5 * cm, 12.3 * cm])
        t.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ("LINEBELOW", (0, 0), (-1, -1), 1, LINE),
        ]))
        return t

    def _make_chip(self, text: str, styles) -> Table:
        """Étiquette arrondie (tag) pour une compétence — largeur ajustée au texte."""
        p = Paragraph(text, styles["cv_chip"])
        text_w = pdfmetrics.stringWidth(text, "Helvetica-Bold", 8.2)
        w = text_w + 20  # marge de sécurité pour éviter tout retour à la ligne
        t = Table([[p]], colWidths=[w])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), TINT_STRONG),
            ("ROUNDEDCORNERS", [7, 7, 7, 7]),
            ("LEFTPADDING", (0, 0), (-1, -1), 9),
            ("RIGHTPADDING", (0, 0), (-1, -1), 9),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        return t, w + 6

    def _pack_chips(self, items: List[str], styles, max_width: float) -> List:
        """Empile les compétences en mini-lignes de tags qui se replient (wrap) comme des badges."""
        rows = []
        current_row, current_widths, current_w = [], [], 0.0
        for item in items:
            chip, w = self._make_chip(item, styles)
            if current_row and current_w + w > max_width:
                rows.append((current_row, current_widths))
                current_row, current_widths, current_w = [], [], 0.0
            current_row.append(chip)
            current_widths.append(w)
            current_w += w
        if current_row:
            rows.append((current_row, current_widths))

        flowables = []
        for row, widths in rows:
            t = Table([row], colWidths=widths)
            t.setStyle(TableStyle([
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]))
            flowables.append(t)
        return flowables

    def _circular_monogram_path(self, initials: str, bg_hex: str = "#0F766E", ring_hex: str = "#CCEEEA") -> Optional[str]:
        """Génère un monogramme circulaire (PNG) pour le CV Word — fiable, sans police emoji."""
        try:
            from PIL import Image as PILImage, ImageDraw, ImageFont
            import tempfile

            size = 240
            ring = 8
            out = PILImage.new("RGBA", (size, size), (0, 0, 0, 0))
            draw = ImageDraw.Draw(out)
            draw.ellipse((0, 0, size, size), fill=ring_hex)
            draw.ellipse((ring, ring, size - ring, size - ring), fill=bg_hex)

            font = None
            for fp in ("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
                       "/usr/local/lib/python3.12/dist-packages/matplotlib/mpl-data/fonts/ttf/DejaVuSans-Bold.ttf"):
                if os.path.exists(fp):
                    font = ImageFont.truetype(fp, 88)
                    break
            if font is None:
                font = ImageFont.load_default()

            text = initials.upper()
            bbox = draw.textbbox((0, 0), text, font=font)
            tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
            draw.text(((size - tw) / 2 - bbox[0], (size - th) / 2 - bbox[1]), text,
                       font=font, fill="#FFFFFF")

            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            out.save(tmp.name)
            return tmp.name
        except Exception as e:
            logger.warning(f"Génération du monogramme circulaire impossible : {e}")
            return None

    def _circular_image_path(self, src_path: str, ring_color_hex: str = "#0F766E") -> Optional[str]:
        """Recadre une photo en cercle avec un fin anneau coloré ; renvoie le chemin du PNG généré."""
        try:
            from PIL import Image as PILImage, ImageDraw, ImageOps
            im = PILImage.open(src_path).convert("RGB")
            size = min(im.size)
            im = ImageOps.fit(im, (size, size), centering=(0.5, 0.35))

            ring = max(6, size // 28)
            canvas_size = size + ring * 2
            mask = PILImage.new("L", (size, size), 0)
            ImageDraw.Draw(mask).ellipse((0, 0, size, size), fill=255)

            out = PILImage.new("RGBA", (canvas_size, canvas_size), (0, 0, 0, 0))
            draw = ImageDraw.Draw(out)
            draw.ellipse((0, 0, canvas_size, canvas_size), fill=ring_color_hex)
            out.paste(im, (ring, ring), mask)

            import tempfile
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            out.save(tmp.name)
            return tmp.name
        except Exception as e:
            logger.warning(f"Recadrage circulaire impossible : {e}")
            return None

    # ────────────────────────────────────────────────────────
    # GÉNÉRATION PDF — CV
    # ────────────────────────────────────────────────────────

    def _draw_sidebar_bg(self, canvas, doc):
        canvas.saveState()
        canvas.setFillColor(SIDEBAR_BG)
        canvas.rect(0, 0, 6.5 * cm, A4[1], fill=True, stroke=False)
        canvas.setStrokeColor(LINE)
        canvas.setLineWidth(1)
        canvas.line(6.5 * cm, 0, 6.5 * cm, A4[1])
        canvas.restoreState()

    def _get_avatar_flowable(self, avatar_url: Optional[str], initials: str) -> Table:
        width = 3.6 * cm
        height = 3.6 * cm

        if avatar_url:
            import requests
            import tempfile
            try:
                r = requests.get(avatar_url, timeout=5)
                if r.status_code == 200:
                    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
                    temp_file.write(r.content)
                    temp_file.close()

                    circ_path = self._circular_image_path(temp_file.name)
                    img_path = circ_path or temp_file.name
                    img = RLImage(img_path, width=width, height=height)
                    img_table = Table([[img]], colWidths=[width], rowHeights=[height])
                    img_table.setStyle(TableStyle([
                        ("LEFTPADDING", (0, 0), (-1, -1), 0),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                        ("TOPPADDING", (0, 0), (-1, -1), 0),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ]))
                    return img_table
            except Exception as e:
                logger.warning(f"Impossible de charger l'avatar: {e}")


        # Monogramme circulaire (vectoriel — fiable, sans dépendance à une police emoji)
        w_pt, h_pt = width, height
        d = Drawing(w_pt, h_pt)
        d.add(Circle(w_pt / 2, h_pt / 2, w_pt / 2, fillColor=TINT_STRONG, strokeColor=None))
        d.add(Circle(w_pt / 2, h_pt / 2, w_pt / 2 - 3, fillColor=PRIMARY, strokeColor=None))
        d.add(String(w_pt / 2, h_pt / 2 - 8, initials.upper(),
                      fontName="Helvetica-Bold", fontSize=22, fillColor=colors.white,
                      textAnchor="middle"))
        wrapper = Table([[d]], colWidths=[w_pt])
        wrapper.setStyle(TableStyle([
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        return wrapper

    def _create_progress_bar(self, level: float, total_width=3.6 * cm) -> Table:
        width_total = total_width
        width_fill = max(width_total * level, 0.18 * cm)
        width_empty = width_total - width_fill

        col_widths = [width_fill, width_empty] if width_empty > 0 else [width_total]
        bar = Table([["", ""]] if width_empty > 0 else [[""]], colWidths=col_widths, rowHeights=[0.16 * cm])
        style = [
            ("BACKGROUND", (0, 0), (0, 0), PRIMARY),
            ("ROUNDEDCORNERS", [3, 3, 3, 3]),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ]
        if width_empty > 0:
            style.append(("BACKGROUND", (1, 0), (1, 0), LINE))
        bar.setStyle(TableStyle(style))
        return bar

    def _timeline_entry(self, title_text: str, period_text: str, meta_text: str,
                         bullets: List[str], styles, content_width: float, marker_color=None) -> List:
        """Construit une ligne [puce, contenu] pour une entrée de timeline (poste, diplôme...)."""
        marker_color = marker_color or PRIMARY
        header = Table(
            [[Paragraph(f"<b>{title_text}</b>", styles["cv_exp_title"]),
              Paragraph(period_text, styles["cv_exp_period"]) if period_text else ""]],
            colWidths=[content_width - 3.4 * cm, 3.4 * cm],
            style=TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ])
        )
        content = [header]
        if meta_text:
            content.append(Paragraph(meta_text, styles["cv_exp_meta"]))
        for b in bullets:
            content.append(Paragraph(f"<font color='#0F766E'>–</font>&nbsp;&nbsp;{b}", styles["cv_exp_bullet"]))
        content.append(Spacer(1, 9))
        return [self._dot(marker_color, 6.5), content]

    def _build_timeline_table(self, profile: UserProfileData, target_job: str, styles) -> List:
        """Construit les sections de la colonne principale (Parcours, Formation, Langues, Centres d'intérêt)."""
        blocks = []
        FULL_WIDTH = 12.8 * cm
        MARKER_COL = 0.7 * cm
        CONTENT_COL = FULL_WIDTH - MARKER_COL

        def section_block(rows):
            t = Table(rows, colWidths=[MARKER_COL, CONTENT_COL])
            t.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ("LINEAFTER", (0, 1), (0, -1), 1.3, LINE),
            ]))
            return t

        # 1. PARCOURS PROFESSIONNEL
        if profile.experiences:
            rows = [[self._square_marker(PRIMARY, 9), Paragraph("PARCOURS PROFESSIONNEL", styles["cv_section_right"])]]
            exps = profile.experiences if isinstance(profile.experiences, list) else [profile.experiences]
            for exp in exps:
                if isinstance(exp, dict):
                    job_title = exp.get("title", exp.get("poste", ""))
                    company   = exp.get("company", exp.get("entreprise", ""))
                    period    = exp.get("period", exp.get("periode", ""))
                    desc      = exp.get("description", exp.get("desc", ""))
                    bullets = self._split_sentences(desc)
                    rows.append(self._timeline_entry(job_title, period, company, bullets, styles, CONTENT_COL))
                else:
                    rows.append(self._timeline_entry(str(exp), "", "", [], styles, CONTENT_COL))
            blocks.append(section_block(rows))
            blocks.append(Spacer(1, 8))

        # 2. FORMATION (regroupée par bloc séparé par une ligne vide : diplôme + établissement)
        if profile.education_level:
            rows = [[self._square_marker(PRIMARY, 9), Paragraph("FORMATION", styles["cv_section_right"])]]
            edu_blocks = [b for b in profile.education_level.split("\n\n") if b.strip()]
            if not edu_blocks:
                edu_blocks = [profile.education_level]
            for block in edu_blocks:
                lines = [l.strip() for l in block.split("\n") if l.strip()]
                if not lines:
                    continue
                degree = lines[0]
                school = " — ".join(lines[1:]) if len(lines) > 1 else ""
                rows.append(self._timeline_entry(degree, "", school, [], styles, CONTENT_COL))
            blocks.append(section_block(rows))
            blocks.append(Spacer(1, 8))

        return blocks

    def _cv_left_column(self, profile: UserProfileData, styles) -> List:
        items = []
        SIDE_WIDTH = 5.2 * cm  # largeur utile de la colonne (hors marges)

        # Initiales monogramme
        initials = ""
        if profile.first_name: initials += profile.first_name[0]
        if profile.last_name: initials += profile.last_name[0]
        if not initials: initials = "CV"

        avatar_flow = self._get_avatar_flowable(profile.avatar_url, initials)
        items.append(avatar_flow)
        items.append(Spacer(1, 14))

        def contact_row(label, value):
            return Table(
                [[Paragraph(label, styles["cv_label"]), Paragraph(value, styles["cv_value"])]],
                colWidths=[1.5 * cm, SIDE_WIDTH - 1.5 * cm],
                style=TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ])
            )

        if profile.location:
            items.append(contact_row("LIEU", profile.location))
        if profile.phone:
            items.append(contact_row("TÉL", profile.phone))
        if profile.email:
            items.append(contact_row("EMAIL", profile.email))
        if profile.linkedin:
            items.append(contact_row("LINKEDIN", profile.linkedin))
        if profile.github:
            items.append(contact_row("GITHUB", profile.github))
        if profile.portfolio:
            items.append(contact_row("PORTFOLIO", profile.portfolio))

        items.append(Spacer(1, 4))
        items.append(HRFlowable(width="100%", thickness=1, color=LINE, spaceAfter=12))

        if profile.skills:
            items.append(Paragraph("COMPÉTENCES", styles["cv_section_left"]))
            items.append(Spacer(1, 2))
            items.extend(self._pack_chips(profile.skills, styles, SIDE_WIDTH))
            items.append(Spacer(1, 8))
            items.append(HRFlowable(width="100%", thickness=1, color=LINE, spaceAfter=12))

        if profile.languages:
            items.append(Paragraph("LANGUES", styles["cv_section_left"]))
            items.append(Spacer(1, 4))

            langs = profile.languages if isinstance(profile.languages, list) else list(profile.languages.items())
            for lang in langs:
                if isinstance(lang, (list, tuple)):
                    label, level_str = lang[0], lang[1]
                else:
                    label, level_str = str(lang), "Moyen"

                lower_lvl = level_str.lower()
                if "maternelle" in lower_lvl or "bilingue" in lower_lvl or "parfait" in lower_lvl:
                    level_val = 1.0
                elif "courant" in lower_lvl or "avance" in lower_lvl or "avancé" in lower_lvl:
                    level_val = 0.8
                elif "intermediaire" in lower_lvl or "intermédiaire" in lower_lvl or "moyen" in lower_lvl:
                    level_val = 0.6
                else:
                    level_val = 0.3

                items.append(Paragraph(label, styles["cv_lang_name"]))
                items.append(Paragraph(level_str, styles["cv_lang_level"]))
                items.append(Spacer(1, 3))
                bar = self._create_progress_bar(level_val, total_width=SIDE_WIDTH)
                bar_wrap = Table([[bar]], colWidths=[SIDE_WIDTH],
                                  style=TableStyle([
                                      ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                                      ("TOPPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
                                  ]))
                items.append(bar_wrap)

            items.append(Spacer(1, 2))
            items.append(HRFlowable(width="100%", thickness=1, color=LINE, spaceAfter=12))

        interests_list = []
        if hasattr(profile, 'interests') and profile.interests:
            if isinstance(profile.interests, list):
                interests_list.extend(profile.interests)
            else:
                interests_list.append(str(profile.interests))

        title_text = "CENTRES D'INTÉRÊT"
        if not interests_list and hasattr(profile, 'certifications') and profile.certifications:
            title_text = "CERTIFICATIONS"
            certs = profile.certifications if isinstance(profile.certifications, list) else [profile.certifications]
            for cert in certs:
                label = cert.get("name", str(cert)) if isinstance(cert, dict) else str(cert)
                interests_list.append(label)

        if interests_list:
            items.append(Paragraph(title_text, styles["cv_section_left"]))
            items.append(Spacer(1, 2))
            items.extend(self._pack_chips(interests_list, styles, SIDE_WIDTH))

        return items

    def _build_cv_pdf(self, profile: UserProfileData, target_job: str = "") -> str:
        file_name = f"cv_{uuid.uuid4().hex[:8]}.pdf"
        file_path = os.path.join(OUTPUT_DIR, file_name)

        doc = SimpleDocTemplate(
            file_path, pagesize=A4,
            leftMargin=0, rightMargin=0,
            topMargin=0, bottomMargin=0,
        )

        styles = self._get_styles()
        story = []

        left_col_flowables = self._cv_left_column(profile, styles)
        right_col_flowables = []
        
        full_name = f"{profile.first_name} {profile.last_name}"
        title_txt = target_job or profile.current_title or "Professionnel"
        
        right_col_flowables.append(Spacer(1, 8))
        right_col_flowables.append(Paragraph(full_name, styles["cv_name"]))
        right_col_flowables.append(Paragraph(title_txt.upper(), styles["cv_subtitle"]))
        right_col_flowables.append(Spacer(1, 6))
        accent_rule = Table([[""]], colWidths=[2.6 * cm], rowHeights=[0.1 * cm],
                             style=TableStyle([("BACKGROUND", (0, 0), (-1, -1), PRIMARY),
                                                ("ROUNDEDCORNERS", [2, 2, 2, 2])]))
        right_col_flowables.append(accent_rule)
        right_col_flowables.append(Spacer(1, 14))

        if profile.bio:
            right_col_flowables.append(self._section_header("PROFIL PROFESSIONNEL", styles))
            right_col_flowables.append(Spacer(1, 6))
            right_col_flowables.append(Paragraph(profile.bio, styles["body_normal"]))
            right_col_flowables.append(Spacer(1, 12))

        right_col_flowables.extend(self._build_timeline_table(profile, target_job, styles))

        body_table = Table(
            [[left_col_flowables, right_col_flowables]],
            colWidths=[6.5 * cm, 14.5 * cm],
            rowHeights=None,
        )
        body_table.setStyle(TableStyle([
            ("VALIGN",      (0,0), (-1,-1), "TOP"),
            ("LEFTPADDING", (0,0), (0,0), 0.5 * cm),
            ("RIGHTPADDING",(0,0), (0,0), 0.5 * cm),
            ("LEFTPADDING", (1,0), (1,0), 0.7 * cm),
            ("RIGHTPADDING",(1,0), (1,0), 1.0 * cm),
            ("TOPPADDING",  (0,0), (-1,-1), 1.0 * cm),
            ("BOTTOMPADDING",(0,0), (-1,-1), 1.0 * cm),
        ]))
        story.append(body_table)

        doc.build(
            story,
            onFirstPage=self._draw_sidebar_bg,
            onLaterPages=self._draw_sidebar_bg
        )
        logger.info(f"PDF CV généré : {file_path}")
        return file_path

    # ────────────────────────────────────────────────────────
    # STYLES REPORTLAB
    # ────────────────────────────────────────────────────────

    def _get_styles(self) -> Dict[str, ParagraphStyle]:
        base = getSampleStyleSheet()
        return {
            # Entêtes lettre
            "date":           ParagraphStyle("date",           fontSize=9.5, textColor=TEXT_GREY, alignment=TA_RIGHT, leading=13),
            "letter_object":  ParagraphStyle("letter_object",  fontSize=11.5, textColor=PRIMARY, fontName="Helvetica-Bold", spaceAfter=4),

            # Corps texte
            "body_normal":    ParagraphStyle("body_normal",    fontSize=9.5, textColor=TEXT_DARK, leading=13.5, spaceAfter=4),
            "body_justify":   ParagraphStyle("body_justify",   fontSize=9.5, textColor=TEXT_DARK, leading=14, alignment=TA_JUSTIFY),

            # CV — bandeau d'en-tête
            "cv_name":        ParagraphStyle("cv_name",        fontSize=25, textColor=INK, fontName="Helvetica-Bold", leading=27),
            "cv_subtitle":    ParagraphStyle("cv_subtitle",    fontSize=12.5, textColor=PRIMARY, fontName="Helvetica-Bold", leading=16, spaceBefore=2),

            # CV — titres de section (colonne principale, fond blanc)
            "cv_section_right": ParagraphStyle("cv_section_right", fontSize=11.5, textColor=INK, fontName="Helvetica-Bold", leading=14, spaceBefore=0, spaceAfter=0),

            # CV — sidebar (colonne latérale, fond teinté)
            "cv_section_left":  ParagraphStyle("cv_section_left",  fontSize=9.5, textColor=PRIMARY, fontName="Helvetica-Bold", leading=12, spaceBefore=0, spaceAfter=5),
            "cv_label":       ParagraphStyle("cv_label",       fontSize=7,   textColor=PRIMARY, fontName="Helvetica-Bold", leading=9),
            "cv_value":       ParagraphStyle("cv_value",       fontSize=8.8, textColor=TEXT_DARK, leading=12),
            "cv_bio":         ParagraphStyle("cv_bio",         fontSize=8.8, textColor=TEXT_DARK, leading=12.5, alignment=TA_LEFT),
            "cv_small":       ParagraphStyle("cv_small",       fontSize=8.5, textColor=TEXT_DARK, leading=12),
            "cv_chip":        ParagraphStyle("cv_chip",        fontSize=8.2, textColor=PRIMARY_DK, fontName="Helvetica-Bold", leading=10.5),

            # CV — entrées de parcours / formation
            "cv_exp_title":   ParagraphStyle("cv_exp_title",   fontSize=10,  textColor=INK, fontName="Helvetica-Bold", leading=13),
            "cv_exp_period":  ParagraphStyle("cv_exp_period",  fontSize=8.3, textColor=PRIMARY, fontName="Helvetica-Bold", alignment=TA_RIGHT, leading=13),
            "cv_exp_meta":    ParagraphStyle("cv_exp_meta",    fontSize=8.7, textColor=TEXT_GREY, fontName="Helvetica-Oblique", leading=11.5, spaceAfter=3),
            "cv_exp_bullet":  ParagraphStyle("cv_exp_bullet",  fontSize=8.7, textColor=TEXT_DARK, leading=12.5, alignment=TA_LEFT, leftIndent=2, spaceAfter=2.5),

            # CV — langues
            "cv_lang_name":   ParagraphStyle("cv_lang_name",   fontSize=9,   textColor=INK, fontName="Helvetica-Bold", leading=12),
            "cv_lang_level":  ParagraphStyle("cv_lang_level",  fontSize=7.8, textColor=TEXT_GREY, leading=10),
        }

    @staticmethod
    def _letter_type_label(letter_type: str) -> str:
        labels = {
            "resignation":           "Lettre de démission",
            "internship_request":    "Demande de stage",
            "follow_up":             "Lettre de relance – Candidature",
            "recommendation_request":"Demande de lettre de recommandation",
        }
        return labels.get(letter_type, "Lettre professionnelle")